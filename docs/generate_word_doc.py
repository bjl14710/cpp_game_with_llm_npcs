#!/usr/bin/env python3
"""
Generate Silmulator User Guide as a Microsoft Word (.docx) document.

Usage:
    python docs/generate_word_doc.py
    # Output: docs/Silmulator_User_Guide.docx

Requirements:
    pip install python-docx
"""

from __future__ import annotations
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Install it with:  pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str) -> None:
    """Set a table cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    doc.add_paragraph(text, style=style)


def add_code_block(doc: Document, code: str) -> None:
    """Add a monospace code block paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    shading_elm = OxmlElement("w:pPr")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F8")
    pPr.append(shd)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"Note: {text}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
    p.paragraph_format.left_indent = Inches(0.3)


def add_table_from_rows(
    doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = hdr
        set_cell_bg(cell, "2D2D5A")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def build_cover(doc: Document) -> None:
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Silmulator")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x5A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("SIL/HIL/CIL Virtual Hardware Simulator").font.size = Pt(18)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("User Guide — Version 1.0.0\n2026-05-09").font.size = Pt(12)

    doc.add_page_break()


def build_toc(doc: Document) -> None:
    add_section_heading(doc, "Table of Contents")
    toc_items = [
        ("1.", "Introduction", 2),
        ("2.", "System Requirements", 2),
        ("3.", "Installation", 2),
        ("4.", "Quick Start", 2),
        ("5.", "Application Layout", 2),
        ("6.", "Virtual Hardware Reference", 2),
        ("  6.1", "Oscilloscope — Tektronix MSO54", 4),
        ("  6.2", "Power Supply — Keysight E36312A", 4),
        ("  6.3", "Function Generator — Keysight 33522B", 4),
        ("  6.4", "Multimeter — Keysight 34401A", 4),
        ("  6.5", "MSP430 Microcontroller", 4),
        ("  6.6", "Code Input Block", 4),
        ("7.", "Building a Simulation", 2),
        ("8.", "SCPI Console", 2),
        ("9.", "Saving and Loading Layouts", 2),
        ("10.", "Keyboard Shortcuts", 2),
        ("11.", "Troubleshooting", 2),
    ]
    for num, title, _level in toc_items:
        p = doc.add_paragraph(f"{num}  {title}", style="Normal")
        p.paragraph_format.left_indent = Inches(0.3 if num.startswith(" ") else 0.0)
    doc.add_page_break()


def build_introduction(doc: Document) -> None:
    add_section_heading(doc, "1. Introduction")
    add_paragraph(doc,
        "Silmulator is a Software-in-the-Loop (SIL) / Hardware-in-the-Loop (HIL) / "
        "Code-in-the-Loop (CIL) virtual laboratory bench. Instead of wiring up physical "
        "instruments, you place virtual blocks on a canvas, connect them with wires, and "
        "interact with them using the same SCPI (Standard Commands for Programmable "
        "Instruments) that your real measurement code already sends."
    )
    add_paragraph(doc,
        "This lets you test and debug firmware and measurement scripts at your desk, "
        "without physical hardware, catching command errors and response-handling bugs "
        "before they ever reach a real instrument."
    )
    add_section_heading(doc, "What it replaces", level=2)
    add_table_from_rows(doc,
        ["Physical Equipment", "Silmulator Virtual Equivalent"],
        [
            ["Tektronix MSO54 oscilloscope", "Oscilloscope block"],
            ["Keysight E36312A triple-output PSU", "Power Supply block"],
            ["Keysight 33522B dual-channel AWG", "Function Generator block"],
            ["Keysight 34401A 6½-digit DMM", "Multimeter block"],
            ["TI MSP430FR5994 microcontroller", "MSP430 block"],
            ["VISA/GPIB cable", "Wire connection on canvas"],
        ],
        col_widths=[3.0, 3.0],
    )
    doc.add_page_break()


def build_requirements(doc: Document) -> None:
    add_section_heading(doc, "2. System Requirements")
    add_section_heading(doc, "Native Installation", level=2)
    add_table_from_rows(doc,
        ["Component", "Minimum Version"],
        [
            ["OS", "macOS 12+, Windows 10, or Ubuntu 20.04+"],
            ["Python", "3.9 or newer"],
            ["RAM", "512 MB"],
            ["Disk", "300 MB"],
            ["Display", "1024 × 768 or larger"],
        ],
        col_widths=[2.0, 4.0],
    )
    add_section_heading(doc, "Docker Installation", level=2)
    add_table_from_rows(doc,
        ["Component", "Requirement"],
        [
            ["Docker Desktop", "4.0 or newer"],
            ["RAM (allocated to Docker)", "1 GB"],
            ["Browser (web UI)", "Chrome, Firefox, Safari, or Edge"],
            ["Disk", "2 GB"],
        ],
        col_widths=[2.5, 3.5],
    )
    doc.add_page_break()


def build_installation(doc: Document) -> None:
    add_section_heading(doc, "3. Installation")

    add_section_heading(doc, "Option A: pip install (recommended)", level=2)
    add_code_block(doc,
        "git clone <repository-url>\n"
        "cd Silmulator\n"
        "python3 -m venv .venv\n"
        "source .venv/bin/activate    # macOS/Linux\n"
        "# .venv\\Scripts\\activate   # Windows\n"
        "pip install -e .\n"
        "silmulator"
    )

    add_section_heading(doc, "Option B: Docker — Browser Access (NoVNC)", level=2)
    add_paragraph(doc,
        "No display setup required. The application streams to your web browser."
    )
    add_code_block(doc,
        "docker compose up -d\n"
        "# Open http://localhost:6080/vnc.html\n"
        "# Password: silmulator"
    )

    add_section_heading(doc, "Option C: Docker — Native Window", level=2)
    add_paragraph(doc,
        "Renders in a native window on your host. Requires XQuartz on macOS or X11 on Linux."
    )
    add_code_block(doc,
        "xhost +localhost   # macOS: allow Docker to connect to XQuartz\n"
        "docker compose -f docker-compose.dev.yml up"
    )

    add_section_heading(doc, "Option D: From Source with C++ Acceleration", level=2)
    add_paragraph(doc,
        "Builds the optional high-performance C++ SCPI parser."
    )
    add_code_block(doc,
        "brew install cmake pybind11   # macOS\n"
        "# apt install cmake python3-dev pybind11-dev g++  # Ubuntu\n"
        "./build_cpp.sh\n"
        "pip install -e .\n"
        "silmulator"
    )
    doc.add_page_break()


def build_quick_start(doc: Document) -> None:
    add_section_heading(doc, "4. Quick Start")
    steps = [
        "Launch the application: silmulator (or python main.py).",
        "In the Library panel on the left, double-click Power Supply to place it on the canvas.",
        "Double-click Multimeter to place a multimeter.",
        "Click the output port (small circle) on the Power Supply block and drag to the input port on the Multimeter.",
        "Right-click the Power Supply → Send SCPI command → type VOLT 5.0 → press Enter.",
        "Right-click the Power Supply → Send SCPI command → type OUTP ON.",
        "Right-click the Multimeter → Send SCPI command → type MEAS:VOLT:DC?",
        "The response +5.00000000E+00 appears in the Console at the bottom.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)
    doc.add_paragraph()
    doc.add_page_break()


def build_hardware_reference(doc: Document) -> None:
    add_section_heading(doc, "6. Virtual Hardware Reference")

    # Oscilloscope
    add_section_heading(doc, "6.1 Oscilloscope — Tektronix MSO54", level=2)
    add_paragraph(doc,
        "4-channel mixed-signal oscilloscope with full SCPI command coverage. "
        "Supports waveform injection, measurements, and CURVE? data transfer."
    )
    add_table_from_rows(doc,
        ["Command", "Description"],
        [
            ["CH1:SCALE 0.5", "500 mV per division"],
            ["HOR:SCALE 1E-3", "1 ms per division"],
            ["TRIG:A:EDGE:SOURCE CH1", "Trigger source channel 1"],
            ["ACQ:STATE 1", "Start acquisition"],
            ["MEASUREMENT:MEAS1:TYPE FREQUENCY", "Configure measurement 1"],
            ["MEASUREMENT:MEAS1:VALUE?", "Query measured frequency"],
            ["CURVE?", "Retrieve waveform data"],
            ["*RST", "Reset to defaults"],
        ],
        col_widths=[2.8, 3.2],
    )

    # Power Supply
    add_section_heading(doc, "6.2 Power Supply — Keysight E36312A", level=2)
    add_paragraph(doc,
        "Three independently controlled DC output channels with OVP/OCP protection simulation."
    )
    add_table_from_rows(doc,
        ["Command", "Description"],
        [
            ["INST:NSEL 1", "Select channel 1"],
            ["VOLT 3.3", "Set voltage to 3.3 V"],
            ["CURR 1.0", "Current limit 1.0 A"],
            ["OUTP ON", "Enable output"],
            ["MEAS:VOLT:DC?", "Measure output voltage"],
            ["MEAS:CURR:DC?", "Measure output current"],
            ["APPL 5.0, 2.0", "Set voltage and current in one command"],
        ],
        col_widths=[2.8, 3.2],
    )

    # Function Generator
    add_section_heading(doc, "6.3 Function Generator — Keysight 33522B", level=2)
    add_paragraph(doc,
        "Dual-channel arbitrary waveform generator. Generates SIN, SQU, RAMP, PULS, NOIS, DC, "
        "USER, and ARB functions with AM/FM/PM/PWM modulation, burst, and sweep."
    )
    add_table_from_rows(doc,
        ["Command", "Description"],
        [
            ["SOURCE1:FUNC SIN", "Sine wave on channel 1"],
            ["SOURCE1:FREQ 1000", "1 kHz frequency"],
            ["SOURCE1:VOLT 2.0", "2 Vpp amplitude"],
            ["OUTPUT1 ON", "Enable channel 1 output"],
            ["SOURCE1:AM:STAT ON", "Enable AM modulation"],
            ["SOURCE1:SWEEP:STAT ON", "Enable frequency sweep"],
        ],
        col_widths=[2.8, 3.2],
    )

    # Multimeter
    add_section_heading(doc, "6.4 Multimeter — Keysight 34401A", level=2)
    add_paragraph(doc,
        "6½-digit DMM supporting DC/AC voltage, DC/AC current, resistance, "
        "continuity, diode, and temperature measurements."
    )
    add_table_from_rows(doc,
        ["Command", "Description"],
        [
            ["MEAS:VOLT:DC?", "DC voltage measurement"],
            ["MEAS:VOLT:AC?", "AC voltage (RMS)"],
            ["MEAS:CURR:DC?", "DC current"],
            ["MEAS:RES?", "2-wire resistance"],
            ["CONF:VOLT:DC 10, 0.001", "Configure range and resolution"],
            ["READ?", "Trigger and return one reading"],
            ["CALC:FUNC LIM; CALC:STAT ON", "Enable limit testing"],
        ],
        col_widths=[2.8, 3.2],
    )

    # MSP430
    add_section_heading(doc, "6.5 MSP430 Microcontroller — TI MSP430FR5994", level=2)
    add_paragraph(doc,
        "Full register-level simulation including GPIO, 12-channel ADC, timers, "
        "UART, SPI, I2C, clocks, and power management modes."
    )
    add_table_from_rows(doc,
        ["Peripheral", "Example Command"],
        [
            ["GPIO", "GPIO:PORT1:DIRECTION 0xFF"],
            ["ADC", "ADC:CHANNEL0:INJECT 1.65"],
            ["UART", "UART:A0:BAUD 115200"],
            ["SPI", "SPI:B0:TRANSFER 0xAB"],
            ["I2C", "I2C:B1:WRITE 0x48, 0x00"],
            ["Clocks", "CLK:DCO:FREQ 16000000"],
            ["Power modes", "PM:MODE LPM0"],
        ],
        col_widths=[2.0, 4.0],
    )

    # Code Input
    add_section_heading(doc, "6.6 Code Input Block", level=2)
    add_paragraph(doc,
        "Analyzes C, C++, or Python source code and extracts SCPI command strings. "
        "Recognized patterns include viWrite(), viQueryf(), inst.write(), inst.query(), "
        "visa_write(), scpi_send(), and ibwrt()."
    )
    doc.add_page_break()


def build_keyboard_shortcuts(doc: Document) -> None:
    add_section_heading(doc, "10. Keyboard Shortcuts")
    add_table_from_rows(doc,
        ["Shortcut", "Action"],
        [
            ["Ctrl+N", "New canvas"],
            ["Ctrl+O", "Open layout"],
            ["Ctrl+S", "Save layout"],
            ["Ctrl+A", "Select all blocks"],
            ["Delete", "Delete selected blocks or wires"],
            ["Ctrl+0", "Reset zoom to 100%"],
            ["Ctrl+Shift+F", "Fit all blocks in view"],
            ["Middle-mouse drag", "Pan canvas"],
            ["Scroll wheel", "Zoom in / out"],
        ],
        col_widths=[2.5, 3.5],
    )
    doc.add_page_break()


def build_troubleshooting(doc: Document) -> None:
    add_section_heading(doc, "11. Troubleshooting")
    issues = [
        (
            "Application does not start",
            "Verify Python ≥ 3.9 is installed. Check that PyQt6 is present: "
            "python3 -c \"import PyQt6\". On Linux, install: apt install libxcb-cursor0."
        ),
        (
            "SCPI Error -113: Undefined header",
            "The command does not match any handler. Check spelling and that the "
            "command belongs to the selected device type."
        ),
        (
            "Docker: browser shows blank screen",
            "Wait 10–15 seconds for the VNC server to start. Refresh the browser."
        ),
        (
            "Docker: VNC password prompt loops",
            "Default password is 'silmulator'. To reset: docker compose down -v && docker compose up -d"
        ),
        (
            "Waveform CURVE? returns zeros",
            "Run ACQ:STATE 1 first, then call inject_waveform() via the Python API "
            "to supply waveform data to the oscilloscope."
        ),
    ]
    for issue, solution in issues:
        p = doc.add_paragraph()
        p.add_run(issue).bold = True
        doc.add_paragraph(solution, style="Normal")
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    out_path = Path(__file__).parent / "Silmulator_User_Guide.docx"

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Heading style colours
    for level, size in [(1, 16), (2, 13), (3, 11)]:
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x5A)
            style.font.size = Pt(size)

    build_cover(doc)
    build_toc(doc)
    build_introduction(doc)
    build_requirements(doc)
    build_installation(doc)
    build_quick_start(doc)

    # Application layout (abbreviated)
    add_section_heading(doc, "5. Application Layout")
    add_paragraph(doc,
        "The application window is divided into three areas: "
        "(1) the Library panel on the left for placing instrument blocks, "
        "(2) the Canvas in the center where blocks are arranged and connected, "
        "and (3) the SCPI Console at the bottom for typing commands and viewing responses."
    )
    doc.add_page_break()

    build_hardware_reference(doc)

    # Building a simulation
    add_section_heading(doc, "7. Building a Simulation")
    steps = [
        "Place hardware blocks on the canvas from the Library panel.",
        "Connect blocks with wires by clicking output ports and dragging to input ports.",
        "Send SCPI commands via the right-click menu, the SCPI Console, or the Code Input block.",
        "Observe responses in the Console panel.",
        "Optionally inject signal data using the Python API.",
        "Save the layout for reuse.",
    ]
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)
    doc.add_page_break()

    # SCPI Console
    add_section_heading(doc, "8. SCPI Console")
    add_paragraph(doc,
        "The console at the bottom of the window accepts SCPI commands typed directly. "
        "Commands are sent to the currently selected block on the canvas. "
        "The console shows '>' for commands you send and '<' for device responses."
    )
    doc.add_page_break()

    # Save and load
    add_section_heading(doc, "9. Saving and Loading Layouts")
    add_paragraph(doc,
        "Layouts are saved as .sil JSON files that preserve block types, positions, "
        "wire connections, per-block device state, and Code Input content."
    )
    add_table_from_rows(doc,
        ["Action", "Shortcut"],
        [
            ["Save", "Ctrl+S"],
            ["Save As", "(menu only)"],
            ["Open", "Ctrl+O"],
            ["New Canvas", "Ctrl+N"],
        ],
        col_widths=[3.0, 3.0],
    )
    doc.add_page_break()

    build_keyboard_shortcuts(doc)
    build_troubleshooting(doc)

    doc.save(str(out_path))
    print(f"Word document saved to: {out_path}")


if __name__ == "__main__":
    main()
